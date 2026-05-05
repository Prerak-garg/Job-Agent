import PyPDF2
from docx import Document
from pathlib import Path


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    if path.suffix.lower() == ".pdf":
        return _extract_pdf(file_path)
    elif path.suffix.lower() == ".docx":
        return _extract_docx(file_path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def _extract_pdf(file_path: str) -> str:
    text = []
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text.append(page.extract_text() or "")
    return "\n".join(text)


def _extract_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def apply_fixes_to_docx(original_path: str, fixes: dict, output_path: str):
    doc = Document(original_path)
    for para in doc.paragraphs:
        for original, replacement in fixes.items():
            if original in para.text:
                for run in para.runs:
                    if original in run.text:
                        run.text = run.text.replace(original, replacement)
    doc.save(output_path)
