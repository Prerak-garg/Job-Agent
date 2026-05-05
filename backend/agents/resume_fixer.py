import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from utils.ollama_client import chat

SYSTEM = """You are a professional resume writer. Rewrite resume sections to be ATS-optimised,
impactful, and keyword-rich. Return ONLY valid JSON."""

PROMPT = """Rewrite the following resume section based on the fix instruction.

Fix instruction: {fix}
Original text: {original}
Full resume context: {context}

Return JSON:
{{
  "rewritten": "<the improved text, same format as original — just the text, no labels>"
}}"""


async def apply_fix(fix: dict, resume_text: str) -> str:
    result = await chat(
        prompt=PROMPT.format(
            fix=fix["suggestion"],
            original=fix.get("original", ""),
            context=resume_text[:2000],
        ),
        model="smart",
        system=SYSTEM,
        json_mode=True,
    )
    try:
        return json.loads(result).get("rewritten", fix["suggestion"])
    except Exception:
        return fix["suggestion"]


async def apply_all_fixes(resume_text: str, approved_fixes: list[dict]) -> str:
    updated = resume_text
    for fix in approved_fixes:
        if fix.get("original") and fix.get("original") in updated:
            rewritten = await apply_fix(fix, resume_text)
            updated = updated.replace(fix["original"], rewritten)
    return updated


def save_fixed_docx(original_path: str, fixed_text: str, output_path: str):
    if Path(original_path).suffix.lower() == ".docx":
        doc = Document(original_path)
        for para in doc.paragraphs:
            if para.text.strip():
                for fix_line in fixed_text.split("\n"):
                    if fix_line.strip() and fix_line.strip() in para.text:
                        for run in para.runs:
                            if run.text:
                                run.text = fix_line.strip()
                                break
        doc.save(output_path)
    else:
        out = Document()
        for line in fixed_text.split("\n"):
            p = out.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)
        out.save(output_path)
